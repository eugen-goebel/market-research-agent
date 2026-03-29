"""Tests for DOCX report generation."""

import os
from docx import Document

from agents.mock_data import SAP_MOCK
from agents.analyst import AnalysisResult, SWOTAnalysis, Competitor
from utils.report_generator import generate_docx_report


class TestReportGeneration:
    """Test the full DOCX generation pipeline."""

    def test_generates_docx_file(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        assert os.path.exists(path)
        assert path.endswith(".docx")

    def test_filename_contains_company(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        assert "sap_se" in os.path.basename(path)

    def test_file_not_empty(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        assert os.path.getsize(path) > 1000  # DOCX should be at least 1KB

    def test_docx_is_valid(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        assert len(doc.paragraphs) > 10

    def test_contains_company_name(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "SAP SE" in all_text

    def test_contains_all_sections(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        expected_sections = [
            "MARKET RESEARCH REPORT",
            "Executive Summary",
            "Company Overview",
            "Market Position",
            "SWOT Analysis",
            "Competitive Landscape",
            "Key Industry Trends",
            "Strategic Outlook",
            "Risk Factors",
        ]
        for section in expected_sections:
            assert section in all_text, f"Missing section: {section}"

    def test_contains_tables(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        # SWOT (2 tables) + competitor table = at least 3
        assert len(doc.tables) >= 3

    def test_competitor_table_has_correct_rows(self, tmp_path):
        path = generate_docx_report("SAP SE", SAP_MOCK, output_dir=str(tmp_path))
        doc = Document(path)
        # The last table should be the competitor table
        comp_table = doc.tables[-1]
        # Header row + data rows
        expected_rows = 1 + len(SAP_MOCK.top_competitors)
        assert len(comp_table.rows) == expected_rows

    def test_creates_output_directory(self, tmp_path):
        new_dir = str(tmp_path / "new_subdir" / "reports")
        path = generate_docx_report("Test Corp", SAP_MOCK, output_dir=new_dir)
        assert os.path.exists(path)
        assert os.path.isdir(new_dir)

    def test_special_characters_in_company_name(self, tmp_path):
        path = generate_docx_report("Müller & Co / GmbH", SAP_MOCK, output_dir=str(tmp_path))
        assert os.path.exists(path)
        assert "müller" in os.path.basename(path)


class TestReportWithMinimalData:
    """Test report generation with edge cases."""

    def _minimal_analysis(self):
        return AnalysisResult(
            executive_summary="Summary.",
            company_overview="Overview.",
            market_position="Position.",
            swot=SWOTAnalysis(
                strengths=["S"], weaknesses=["W"],
                opportunities=["O"], threats=["T"],
            ),
            top_competitors=[
                Competitor(
                    name="Rival",
                    overview="A rival company.",
                    key_strength="Fast",
                    key_weakness="Expensive",
                )
            ],
            key_trends=["Trend 1"],
            investment_thesis="Thesis.",
            risk_factors=["Risk 1"],
        )

    def test_minimal_data_generates(self, tmp_path):
        path = generate_docx_report(
            "Minimal Corp", self._minimal_analysis(), output_dir=str(tmp_path)
        )
        assert os.path.exists(path)

    def test_single_competitor(self, tmp_path):
        path = generate_docx_report(
            "Single Comp", self._minimal_analysis(), output_dir=str(tmp_path)
        )
        doc = Document(path)
        comp_table = doc.tables[-1]
        assert len(comp_table.rows) == 2  # header + 1 competitor
