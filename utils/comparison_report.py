"""
Comparison Report Generator — Side-by-side competitor comparison DOCX report.

Takes multiple AnalysisResult objects and produces a single report with
SWOT comparison matrices, competitor overlap tables, and trend analysis.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.analyst import AnalysisResult


# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
COLOR_DARK = RGBColor(0x1A, 0x37, 0x6C)
COLOR_ACCENT = RGBColor(0x2E, 0x6D, 0xB4)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_BG = "F0F5FF"
COLOR_ALT_BG = "FFFFFF"

COMPANY_COLORS = ["1A6C3C", "1A376C", "6C1A3C"]

SWOT_LABELS = [
    ("STRENGTHS", "1D7A4A", "F0FFF4"),
    ("WEAKNESSES", "92600D", "FFFBF0"),
    ("OPPORTUNITIES", "0D666B", "F0FFFF"),
    ("THREATS", "9B1C1C", "FFF5F5"),
]


def _set_bg(cell, color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def _heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.font.color.rgb = COLOR_DARK
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_ACCENT
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)


def _hr(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DB4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_cover(doc, companies):
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("COMPETITOR COMPARISON REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(" vs ".join(companies))
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = COLOR_ACCENT

    doc.add_paragraph()
    doc.add_paragraph()

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = d.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

    doc.add_page_break()


def _text_comparison(doc, companies, analyses, field):
    """Side-by-side table for a text field."""
    cols = len(companies)
    table = doc.add_table(rows=2, cols=cols)
    table.style = "Table Grid"
    table.autofit = False

    w = Inches(6.0 / cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = w

    for i, name in enumerate(companies):
        cell = table.rows[0].cells[i]
        _set_bg(cell, COMPANY_COLORS[i % len(COMPANY_COLORS)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_WHITE

    for i, a in enumerate(analyses):
        cell = table.rows[1].cells[i]
        _set_bg(cell, COLOR_LIGHT_BG if i % 2 == 0 else COLOR_ALT_BG)
        r = cell.paragraphs[0].add_run(getattr(a, field, ""))
        r.font.size = Pt(9.5)

    doc.add_paragraph()


def _list_comparison(doc, companies, analyses, field):
    """Side-by-side table for a list field."""
    cols = len(companies)
    table = doc.add_table(rows=2, cols=cols)
    table.style = "Table Grid"
    table.autofit = False

    w = Inches(6.0 / cols)
    for col in table.columns:
        for cell in col.cells:
            cell.width = w

    for i, name in enumerate(companies):
        cell = table.rows[0].cells[i]
        _set_bg(cell, COMPANY_COLORS[i % len(COMPANY_COLORS)])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(name)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_WHITE

    for i, a in enumerate(analyses):
        cell = table.rows[1].cells[i]
        _set_bg(cell, COLOR_LIGHT_BG if i % 2 == 0 else COLOR_ALT_BG)
        for item in getattr(a, field, []):
            para = cell.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.15)
            r = para.add_run(item)
            r.font.size = Pt(9)

    doc.add_paragraph()


def _swot_comparison(doc, companies, analyses):
    """SWOT matrix comparison: one row per SWOT category, columns per company."""
    swot_fields = ["strengths", "weaknesses", "opportunities", "threats"]

    for (label, header_color, bg_color), field in zip(SWOT_LABELS, swot_fields):
        cols = len(companies) + 1
        table = doc.add_table(rows=2, cols=cols)
        table.style = "Table Grid"
        table.autofit = False

        label_w = Inches(1.2)
        col_w = Inches(4.8 / len(companies))

        table.rows[0].cells[0].width = label_w
        for i in range(1, cols):
            table.rows[0].cells[i].width = col_w

        # Header row
        cell0 = table.rows[0].cells[0]
        _set_bg(cell0, header_color)
        p = cell0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_WHITE

        for i, name in enumerate(companies):
            cell = table.rows[0].cells[i + 1]
            _set_bg(cell, COMPANY_COLORS[i % len(COMPANY_COLORS)])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(name)
            r.font.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_WHITE

        # Content row
        label_cell = table.rows[1].cells[0]
        label_cell.width = label_w
        _set_bg(label_cell, bg_color)

        for i, a in enumerate(analyses):
            cell = table.rows[1].cells[i + 1]
            cell.width = col_w
            _set_bg(cell, bg_color)
            items = getattr(a.swot, field, [])
            for item in items:
                para = cell.add_paragraph(style="List Bullet")
                para.paragraph_format.left_indent = Inches(0.1)
                r = para.add_run(item)
                r.font.size = Pt(9)

        doc.add_paragraph()


def _competitor_table(doc, companies, analyses):
    """Combined competitor landscape table."""
    headers = ["Analyzed Company", "Competitor", "Key Strength", "Key Weakness"]
    rows_data = []
    for name, a in zip(companies, analyses):
        for comp in a.top_competitors[:3]:
            rows_data.append((name, comp.name, comp.key_strength, comp.key_weakness))

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.3), Inches(1.3), Inches(1.8), Inches(1.8)]
    for i, w in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = w

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_bg(cell, "1A376C")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_WHITE

    for idx, (company, comp_name, strength, weakness) in enumerate(rows_data):
        row = table.rows[idx + 1]
        bg = COLOR_LIGHT_BG if idx % 2 == 0 else COLOR_ALT_BG
        for ci, val in enumerate([company, comp_name, strength, weakness]):
            cell = row.cells[ci]
            _set_bg(cell, bg)
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9)
            if ci <= 1:
                r.font.bold = True

    doc.add_paragraph()


def generate_comparison_report(
    companies: list[str],
    analyses: list[AnalysisResult],
    output_dir: str = "output",
) -> str:
    """
    Generate a side-by-side competitor comparison DOCX report.

    Args:
        companies: List of company names
        analyses:  Corresponding list of AnalysisResult objects
        output_dir: Output directory

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    _build_cover(doc, companies)

    _heading(doc, "Executive Summary Comparison")
    _hr(doc)
    _text_comparison(doc, companies, analyses, "executive_summary")

    _heading(doc, "Company Overviews")
    _hr(doc)
    _text_comparison(doc, companies, analyses, "company_overview")

    _heading(doc, "Market Position")
    _hr(doc)
    _text_comparison(doc, companies, analyses, "market_position")

    _heading(doc, "SWOT Comparison")
    _hr(doc)
    _swot_comparison(doc, companies, analyses)

    _heading(doc, "Competitive Landscape")
    _hr(doc)
    _competitor_table(doc, companies, analyses)

    _heading(doc, "Key Industry Trends")
    _hr(doc)
    _list_comparison(doc, companies, analyses, "key_trends")

    _heading(doc, "Strategic Outlook")
    _hr(doc)
    _text_comparison(doc, companies, analyses, "investment_thesis")

    _heading(doc, "Risk Factors")
    _hr(doc)
    _list_comparison(doc, companies, analyses, "risk_factors")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Generated by Multi-Agent Market Research System  •  {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    safe = "_vs_".join(c.lower().replace(" ", "_")[:20] for c in companies)
    filename = f"competitor_comparison_{safe}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
