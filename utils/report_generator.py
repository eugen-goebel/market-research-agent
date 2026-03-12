"""
Report Generator — Produces a professional DOCX market research report.

Transforms the structured AnalysisResult from the AnalysisAgent into a
formatted Word document with cover page, SWOT table, competitor matrix, etc.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from agents.analyst import AnalysisResult


# ---------------------------------------------------------------------------
# Color palette (corporate blue theme)
# ---------------------------------------------------------------------------
COLOR_DARK_BLUE = RGBColor(0x1A, 0x37, 0x6C)   # #1A376C — headers
COLOR_MED_BLUE  = RGBColor(0x2E, 0x6D, 0xB4)   # #2E6DB4 — accents
COLOR_LIGHT_BG  = RGBColor(0xF0, 0xF5, 0xFF)   # #F0F5FF — table backgrounds
COLOR_GREEN     = RGBColor(0x1D, 0x7A, 0x4A)   # SWOT strengths
COLOR_RED       = RGBColor(0x9B, 0x1C, 0x1C)   # SWOT threats
COLOR_AMBER     = RGBColor(0x92, 0x60, 0x0D)   # SWOT weaknesses
COLOR_TEAL      = RGBColor(0x0D, 0x66, 0x6B)   # SWOT opportunities
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_LIGHT_GRAY = RGBColor(0xF7, 0xF8, 0xFA)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_background(cell, hex_color: str):
    """Set a table cell's background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a styled section heading."""
    para = doc.add_heading(text, level=level)
    run = para.runs[0]
    run.font.color.rgb = COLOR_DARK_BLUE
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = COLOR_MED_BLUE
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    return para


def _add_bullet_list(doc: Document, items: list[str], indent: int = 0):
    """Add a bulleted list of items."""
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
        run = para.add_run(item)
        run.font.size = Pt(10.5)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal divider."""
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E6DB4")
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_cover_page(doc: Document, company: str):
    """Create a styled cover page."""
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MARKET RESEARCH REPORT")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK_BLUE

    doc.add_paragraph()

    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = company_para.add_run(company)
    run2.font.size = Pt(22)
    run2.font.bold = True
    run2.font.color.rgb = COLOR_MED_BLUE

    doc.add_paragraph()
    doc.add_paragraph()

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    run3.font.size = Pt(11)
    run3.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run3.font.italic = True

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = sub_para.add_run("AI-Powered Competitive Intelligence  •  Multi-Agent Research System")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    run4.font.italic = True

    doc.add_page_break()


def _build_swot_table(doc: Document, swot):
    """Build a 2×2 SWOT matrix table."""
    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    col_width = Inches(3.1)
    for col in table.columns:
        for cell in col.cells:
            cell.width = col_width

    # Header row
    headers = [("STRENGTHS", "1D7A4A"), ("WEAKNESSES", "92600D"),
               ("OPPORTUNITIES", "0D666B"), ("THREATS", "9B1C1C")]

    header_row = table.rows[0]
    for i, (label, color) in enumerate(headers[:2]):
        cell = header_row.cells[i]
        _set_cell_background(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_WHITE

    # Content rows
    content_row_0 = table.rows[1]
    for i, items in enumerate([swot.strengths, swot.weaknesses]):
        cell = content_row_0.cells[i]
        _set_cell_background(cell, "F0FFF4" if i == 0 else "FFFBF0")
        for item in items:
            para = cell.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.15)
            run = para.add_run(item)
            run.font.size = Pt(10)

    # Opportunities / Threats — second table for cleaner layout
    doc.add_paragraph()
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = "Table Grid"
    table2.autofit = False

    header_row3 = table2.rows[0]
    for i, (label, color) in enumerate(headers[2:]):
        cell = header_row3.cells[i]
        _set_cell_background(cell, color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(label)
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOR_WHITE
        for col in table2.columns:
            for c in col.cells:
                c.width = col_width

    content_row_1 = table2.rows[1]
    for i, items in enumerate([swot.opportunities, swot.threats]):
        cell = content_row_1.cells[i]
        _set_cell_background(cell, "F0FFFF" if i == 0 else "FFF5F5")
        for item in items:
            para = cell.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.15)
            run = para.add_run(item)
            run.font.size = Pt(10)

    # Remove the extra row from the first table
    table._tbl.remove(table.rows[2]._tr)


def _build_competitor_table(doc: Document, competitors):
    """Build a structured competitor comparison table."""
    headers = ["Company", "Overview", "Key Strength", "Key Weakness"]
    table = doc.add_table(rows=1 + len(competitors), cols=4)
    table.style = "Table Grid"
    table.autofit = False

    widths = [Inches(1.2), Inches(2.3), Inches(1.5), Inches(1.5)]
    for i, width in enumerate(widths):
        for row in table.rows:
            row.cells[i].width = width

    # Header row
    header_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = header_row.cells[i]
        _set_cell_background(cell, "1A376C")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    # Data rows
    for row_idx, comp in enumerate(competitors):
        row = table.rows[row_idx + 1]
        bg = "F0F5FF" if row_idx % 2 == 0 else "FFFFFF"

        data = [comp.name, comp.overview, comp.key_strength, comp.key_weakness]
        for col_idx, value in enumerate(data):
            cell = row.cells[col_idx]
            _set_cell_background(cell, bg)
            para = cell.paragraphs[0]
            run = para.add_run(value)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True


# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def generate_docx_report(
    company: str,
    analysis: AnalysisResult,
    output_dir: str = "output",
) -> str:
    """
    Generate a professional DOCX market research report.

    Args:
        company:    Company name used as title
        analysis:   Structured AnalysisResult from the AnalysisAgent
        output_dir: Directory where the file will be saved

    Returns:
        Absolute path to the generated .docx file
    """
    os.makedirs(output_dir, exist_ok=True)

    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    # -----------------------------------------------------------------------
    # 1. Cover Page
    # -----------------------------------------------------------------------
    _build_cover_page(doc, company)

    # -----------------------------------------------------------------------
    # 2. Executive Summary
    # -----------------------------------------------------------------------
    _add_heading(doc, "Executive Summary")
    _add_horizontal_rule(doc)
    para = doc.add_paragraph(analysis.executive_summary)
    para.paragraph_format.space_after = Pt(12)

    # -----------------------------------------------------------------------
    # 3. Company Overview
    # -----------------------------------------------------------------------
    _add_heading(doc, "Company Overview")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.company_overview)

    # -----------------------------------------------------------------------
    # 4. Market Position
    # -----------------------------------------------------------------------
    _add_heading(doc, "Market Position")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.market_position)

    # -----------------------------------------------------------------------
    # 5. SWOT Analysis
    # -----------------------------------------------------------------------
    _add_heading(doc, "SWOT Analysis")
    _add_horizontal_rule(doc)
    _build_swot_table(doc, analysis.swot)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 6. Competitive Landscape
    # -----------------------------------------------------------------------
    _add_heading(doc, "Competitive Landscape")
    _add_horizontal_rule(doc)
    _build_competitor_table(doc, analysis.top_competitors)

    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # 7. Key Industry Trends
    # -----------------------------------------------------------------------
    _add_heading(doc, "Key Industry Trends")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, analysis.key_trends)

    # -----------------------------------------------------------------------
    # 8. Strategic Outlook
    # -----------------------------------------------------------------------
    _add_heading(doc, "Strategic Outlook")
    _add_horizontal_rule(doc)
    doc.add_paragraph(analysis.investment_thesis)

    # -----------------------------------------------------------------------
    # 9. Risk Factors
    # -----------------------------------------------------------------------
    _add_heading(doc, "Risk Factors")
    _add_horizontal_rule(doc)
    _add_bullet_list(doc, analysis.risk_factors)

    # -----------------------------------------------------------------------
    # Footer
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated by Multi-Agent Market Research System  •  {datetime.now().strftime('%Y-%m-%d')}"
    )
    run.font.size = Pt(8.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    safe_name = company.lower().replace(" ", "_").replace("/", "-")
    filename = f"market_research_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)

    return os.path.abspath(filepath)
